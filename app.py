import streamlit as st
from docx import Document
from datetime import datetime
import os
import tempfile

# File path
TEMPLATE_PATH = r"C:\AppSynergyies\AppSynergies\PDF Generator - AS\NDA Template.docx"

def replace_placeholders(template_path, output_path, replacements):
    """Replace placeholders in the Word document."""
    doc = Document(template_path)

    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                for run in para.runs:
                    run.text = run.text.replace(key, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save to temporary file to avoid permission issues
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_output_path = tmp_file.name
    doc.save(temp_output_path)

    return temp_output_path

# Streamlit UI
st.title("NDA Generator")

# Date Selection
selected_date = st.date_input("Select Date", value=datetime.today())
formatted_date = selected_date.strftime("%Y-%m-%d")  # Convert date to string format

# User input fields
client_name = st.text_input("Enter Client Name")
client_address = st.text_area("Enter Client Address")
signature_name = st.text_input("Enter Signature Name")

# Generate Contract
if st.button("Generate Contract"):
    if client_name and client_address and signature_name:
        replacements = {
            "{{Date}}": formatted_date,
            "{{ClientName}}": client_name,
            "{{ClientAddress}}": client_address,
            "{{SignatureName}}": signature_name,
        }

        file_path = replace_placeholders(TEMPLATE_PATH, None, replacements)
        st.success("Contract Generated Successfully!")

        # Provide file download button
        with open(file_path, "rb") as file:
            st.download_button("Download Contract", file, file_name="Filled_Contract.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Clean up temp file
        os.remove(file_path)
    else:
        st.warning("Please fill all fields!")
