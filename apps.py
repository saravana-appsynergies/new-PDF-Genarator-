import streamlit as st
from docx import Document
from datetime import datetime
import os
import tempfile

# Define available templates and fields
TEMPLATES = {
    "Contract Template": {
        "path": "Contract Template.docx",
        "fields": {
            "{{Date}}": {"label": "Select Date", "type": "date"},
            "{{ClientName}}": {"label": "Enter Client Name", "type": "text"},
            "{{ClientAddress}}": {"label": "Enter Client Address", "type": "textarea"},
            "{{SignatureName}}": {"label": "Enter Signature Name", "type": "text"},
            "{{LeaseTerm}}": {"label": "Enter Lease Term", "type": "text"},
        },
    },
    "NDA Template": {
        "path": "NDA Template.docx",
        "fields": {
            "{{Date}}": {"label": "Select Date", "type": "date"},
            "{{ClientName}}": {"label": "Enter Client Name", "type": "text"},
            "{{ClientAddress}}": {"label": "Enter Client Address", "type": "textarea"},
            "{{SignatureName}}": {"label": "Enter Signature Name", "type": "text"},
        },
    },
    "1 Payment Invoice": {
        "path": "Invoice Template - INR - 1 Payment.docx",
        "fields": {
            "{{InvoiceNumber}}": {"label": "Invoice Number", "type": "text"},
            "{{InvoiceDate}}": {"label": "Invoice Date", "type": "date"},
            "{{ClientName}}": {"label": "Client Name", "type": "text"},
            "{{ClientAddress}}": {"label": "Client Address", "type": "textarea"},
            "{{CompanyGST}}": {"label": "Company GST", "type": "text"},
            "{{ProjectName}}": {"label": "Project Name", "type": "text"},
            "{{MobileNumber}}": {"label": "Mobile No (+91-)", "type": "text"},
        },
    },
    "3 EMI Payment Invoice": {
        "path": "Invoice Template - INR - 3 EMI Payment Schedule.docx",
        "fields": {
            "{{InvoiceNumber}}": {"label": "Invoice Number", "type": "text"},
            "{{InvoiceDate}}": {"label": "Invoice Date", "type": "date"},
            "{{ClientName}}": {"label": "Client Name", "type": "text"},
            "{{ClientAddress}}": {"label": "Client Address", "type": "textarea"},
            "{{CompanyGST}}": {"label": "Company GST", "type": "text"},
            "{{ProjectName}}": {"label": "Project Name", "type": "text"},
            "{{MobileNumber}}": {"label": "Mobile No (+91-)", "type": "text"},
        },
    },
    "5 EMI Payment Invoice": {
        "path": "Invoice Template - USD - 5 EMI Payment Schedule.docx",
        "fields": {
            "{{InvoiceNumber}}": {"label": "Invoice Number", "type": "text"},
            "{{InvoiceDate}}": {"label": "Invoice Date", "type": "date"},
            "{{ClientName}}": {"label": "Client Name", "type": "text"},
            "{{ClientAddress}}": {"label": "Client Address", "type": "textarea"},
            "{{CompanyGST}}": {"label": "Company GST", "type": "text"},
            "{{ProjectName}}": {"label": "Project Name", "type": "text"},
            "{{MobileNumber}}": {"label": "Mobile No (+91-)", "type": "text"},
        },
    },
}

def replace_placeholders(template_path, replacements):
    """Replaces placeholders in the Word document template and makes the replaced text bold."""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    
    # Replace placeholders in paragraphs while preserving formatting
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                full_text = para.text
                full_text = full_text.replace(key, value)
                para.clear()  # Clear the paragraph text
                run = para.add_run(full_text)  # Add the modified text
                run.bold = True  # Make the text bold

    # Replace placeholders in tables while preserving formatting
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        full_text = cell.text
                        full_text = full_text.replace(key, value)
                        cell.text = full_text  # Set updated text to the cell
                        # Make the text bold
                        cell.paragraphs[0].runs[0].bold = True

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_output_path = tmp_file.name
        doc.save(temp_output_path)

    return temp_output_path

def extract_placeholders(doc_path):
    """Extracts placeholders dynamically from a Word document."""
    doc = Document(doc_path)
    placeholders = set()

    for para in doc.paragraphs:
        for word in para.text.split():
            if word.startswith("{{") and word.endswith("}}"):
                placeholders.add(word)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for word in cell.text.split():
                    if word.startswith("{{") and word.endswith("}}"):
                        placeholders.add(word)

    return sorted(placeholders)

def get_doc_text(doc_path):
    """Extracts text from a Word document to preview the output."""
    doc = Document(doc_path)
    return "\n".join([para.text for para in doc.paragraphs])

# --- Streamlit UI ---
st.title("📄 Dynamic Document Generator")

# Upload custom template (Optional)
uploaded_template = st.file_uploader("Upload a custom Word template (.docx)", type=["docx"])

if uploaded_template:
    template_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    with open(template_path, "wb") as f:
        f.write(uploaded_template.read())

    # Automatically detect placeholders
    detected_placeholders = extract_placeholders(template_path)

    if not detected_placeholders:
        st.error("❌ No placeholders detected! Please ensure the document contains placeholders like {{Name}}.")

    # Convert detected placeholders into fields
    custom_fields = {ph: {"label": f"Enter value for {ph}", "type": "text"} for ph in detected_placeholders}
    selected_template = {"path": template_path, "fields": custom_fields}

else:
    # Predefined template selection
    template_choice = st.selectbox("Select a Template", list(TEMPLATES.keys()))
    selected_template = TEMPLATES[template_choice]
    template_path = selected_template["path"]

replacements = {}
all_filled = True

st.header("✍ Enter the required details:")

for placeholder, field_info in selected_template["fields"].items():
    if field_info["type"] == "date":
        date_value = st.date_input(field_info["label"], value=datetime.today(), key=placeholder)
        replacements[placeholder] = date_value.strftime("%Y-%m-%d")

    elif field_info["type"] == "text":
        text_value = st.text_input(field_info["label"], key=placeholder)
        if not text_value or text_value.strip() == "":
            all_filled = False
        replacements[placeholder] = text_value.strip()

    elif field_info["type"] == "textarea":
        text_value = st.text_area(field_info["label"], key=placeholder)
        if not text_value or text_value.strip() == "":
            all_filled = False
        replacements[placeholder] = text_value.strip()

st.write("🔍 Debug Info:", replacements)  # Debugging: Ensure all inputs are captured

if st.button("🚀 Generate Document"):
    if all_filled:
        try:
            output_file_path = replace_placeholders(template_path, replacements)
            st.success("✅ Document Generated Successfully!")

            # Preview document content
            preview_text = get_doc_text(output_file_path)
            st.text_area("📄 Preview of Generated Document:", value=preview_text, height=200)

            # Generate dynamic file name
            client_name_clean = replacements.get("{{ClientName}}", "Client").replace(" ", "_")
            dynamic_file_name = f"{client_name_clean}_{template_choice}.docx" if not uploaded_template else "Custom_Document.docx"

            # Provide download button
            with open(output_file_path, "rb") as file:
                st.download_button(
                    "⬇ Download Document",
                    file,
                    file_name=dynamic_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # Cleanup: Remove the temp file after download
            os.remove(output_file_path)

        except Exception as e:
            st.error(f"❌ An error occurred: {e}")

    else:
        st.warning("⚠ Please fill all fields!")

